"""
E-commerce Analytics Report Generator
Created by: Vijeta Thakur (2025)
Generates professional PDF and Word reports with analytics visualizations.
"""

from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import pandas as pd
import plotly.io as pio
import os
from datetime import datetime

class ReportGenerator:
    def __init__(self, df):
        self.df = df
        self.platform_stats = df.groupby('Platform').agg({
            'Order Value (INR)': 'mean',
            'Delivery Time (Minutes)': 'mean',
            'Order ID': 'count'
        }).round(2)
        self.platform_stats.columns = ['Avg Order Value', 'Avg Delivery Time', 'Total Orders']
        
        self.category_stats = df.groupby('Product Category')['Delivery Time (Minutes)'].agg(['mean', 'count']).round(2)
        
        # Ensure output directories exist
        for dir_path in ['output/reports', 'output/temp']:
            if not os.path.exists(dir_path):
                os.makedirs(dir_path)

    def save_figure_as_image(self, fig, filename):
        """Save a plotly figure as an image using kaleido"""
        temp_path = os.path.join('output/temp', filename)
        fig.write_image(temp_path, format='png', engine='kaleido')
        return temp_path

    def save_current_figures(self, figures):
        """Save all current figures as temporary images for the report"""
        self.image_paths = []
        for name, fig in figures.items():
            img_path = self.save_figure_as_image(fig, f'{name}.png')
            self.image_paths.append(img_path)

    def generate_pdf(self, figures):
        """Generate PDF report"""
        # First save all figures
        self.save_current_figures(figures)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f'output/reports/ecommerce_analysis_report_{timestamp}.pdf'
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title and Author
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30
        )
        author_style = ParagraphStyle(
            'Author',
            parent=styles['Normal'],
            fontSize=14,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        story.append(Paragraph('E-commerce Delivery Analytics Report', title_style))
        story.append(Paragraph('Created by: Vijeta Thakur', author_style))
        story.append(Spacer(1, 12))

        # Dataset Overview
        story.append(Paragraph('Dataset Overview', styles['Heading2']))
        overview_data = [
            ['Total Orders', f"{len(self.df):,}"],
            ['Number of Platforms', f"{len(self.df.Platform.unique())}"],
            ['Number of Product Categories', f"{len(self.df['Product Category'].unique())}"]
        ]
        overview_table = Table(overview_data)
        overview_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(overview_table)
        story.append(Spacer(1, 20))

        # Platform Performance
        story.append(Paragraph('Platform Performance Metrics', styles['Heading2']))
        platform_data = [['Platform'] + list(self.platform_stats.columns)]
        for platform in self.platform_stats.index:
            platform_data.append([platform] + [str(self.platform_stats.loc[platform][col]) 
                                            for col in self.platform_stats.columns])
        
        platform_table = Table(platform_data)
        platform_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(platform_table)
        story.append(Spacer(1, 20))

        # Add visualizations
        story.append(Paragraph('Visualizations', styles['Heading2']))
        for img_path in self.image_paths:
            img = Image(img_path, width=450, height=300)
            story.append(img)
            story.append(Spacer(1, 20))

        doc.build(story)
        return filename

    def generate_word(self, figures):
        """Generate Word report"""
        # First save all figures
        self.save_current_figures(figures)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f'output/reports/ecommerce_analysis_report_{timestamp}.docx'
        doc = Document()

        # Title
        doc.add_heading('E-commerce Delivery Analytics Report', 0)
        
        # Author
        author = doc.add_paragraph('Created by: Vijeta Thakur')
        author.alignment = 1  # Center alignment
        author.style = 'Subtitle'

        # Dataset Overview
        doc.add_heading('Dataset Overview', level=1)
        overview_table = doc.add_table(rows=1, cols=2)
        overview_table.style = 'Table Grid'
        overview_table.rows[0].cells[0].text = 'Metric'
        overview_table.rows[0].cells[1].text = 'Value'
        
        row_cells = overview_table.add_row().cells
        row_cells[0].text = 'Total Orders'
        row_cells[1].text = f"{len(self.df):,}"
        
        row_cells = overview_table.add_row().cells
        row_cells[0].text = 'Number of Platforms'
        row_cells[1].text = str(len(self.df.Platform.unique()))
        
        row_cells = overview_table.add_row().cells
        row_cells[0].text = 'Number of Product Categories'
        row_cells[1].text = str(len(self.df['Product Category'].unique()))

        doc.add_paragraph('')

        # Platform Performance
        doc.add_heading('Platform Performance Metrics', level=1)
        platform_table = doc.add_table(rows=1, cols=4)
        platform_table.style = 'Table Grid'
        
        # Header row
        header_cells = platform_table.rows[0].cells
        header_cells[0].text = 'Platform'
        for idx, col in enumerate(self.platform_stats.columns, 1):
            header_cells[idx].text = col

        # Data rows
        for platform in self.platform_stats.index:
            row_cells = platform_table.add_row().cells
            row_cells[0].text = platform
            for idx, col in enumerate(self.platform_stats.columns, 1):
                row_cells[idx].text = str(self.platform_stats.loc[platform][col])

        doc.add_paragraph('')

        # Add visualizations
        doc.add_heading('Visualizations', level=1)
        for img_path in self.image_paths:
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph('')

        doc.save(filename)
        return filename
